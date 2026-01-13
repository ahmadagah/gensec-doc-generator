"""
Document generator for the gensec-template CLI tool.

This module provides functionality to generate lab template documents
in both .docx (Microsoft Word) and .md (Markdown) formats from Lab
data structures.
"""

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

if TYPE_CHECKING:
    from docx.styles.style import _ParagraphStyle

from .models import Lab


def setup_styles(doc: Document) -> None:
    """
    Configure document styles for consistent formatting.

    Sets up the following styles:
    - Heading 1: 24pt, Bold
    - Heading 2: 18pt, Bold
    - Normal: 11pt

    Args:
        doc: The Document object to configure styles for.
    """
    styles = doc.styles

    # Configure Heading 1 style
    heading1_style: "_ParagraphStyle" = styles["Heading 1"]
    heading1_font = heading1_style.font
    heading1_font.size = Pt(24)
    heading1_font.bold = True

    # Configure Heading 2 style
    heading2_style: "_ParagraphStyle" = styles["Heading 2"]
    heading2_font = heading2_style.font
    heading2_font.size = Pt(18)
    heading2_font.bold = True

    # Configure Normal style
    normal_style: "_ParagraphStyle" = styles["Normal"]
    normal_font = normal_style.font
    normal_font.size = Pt(11)


def add_table_of_contents(doc: Document) -> None:
    """
    Add a table of contents placeholder to the document.

    Adds a simple TOC placeholder that Google Docs or Microsoft Word
    can use to auto-generate a table of contents.

    Args:
        doc: The Document object to add the TOC to.
    """
    # Add TOC heading
    toc_heading = doc.add_paragraph("Table of Contents")
    toc_heading.style = doc.styles["Heading 1"]

    # Add placeholder text for TOC
    toc_placeholder = doc.add_paragraph(
        "[Table of Contents - Update in Word/Google Docs to auto-generate]"
    )
    toc_placeholder.italic = True

    # Add page break after TOC
    doc.add_page_break()


def generate_docx(lab: Lab, output_path: str) -> str:
    """
    Generate a .docx template document for the given lab.

    Creates a Word document with:
    - Lab title as Heading 1
    - Section titles as Heading 2 (only sections with questions)
    - Bullet points for each question/task

    Args:
        lab: The Lab object containing all lab data.
        output_path: The path where the document should be saved.

    Returns:
        The output path where the document was saved.
    """
    doc = Document()

    # Setup document styles
    setup_styles(doc)

    # Add lab title as Heading 1
    title = f"{lab.number}: {lab.title}"
    doc.add_heading(title, level=1)

    # Process each section (only include sections with questions)
    for section in lab.sections:
        # Skip sections with no questions
        if not section.questions:
            continue

        # Add section title as Heading 2
        section_title = f"{section.number}. {section.title}"
        doc.add_heading(section_title, level=2)

        # Add bullet points for each question
        for question in section.questions:
            doc.add_paragraph(question.text, style="List Bullet")

        # Add spacing between sections
        doc.add_paragraph()

    # Ensure output directory exists
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    # Save the document
    doc.save(output_path)

    return output_path


def generate_markdown(lab: Lab, output_path: str) -> str:
    """
    Generate a Markdown template document for the given lab.

    Creates a Markdown file with:
    - Lab title as H1 (# heading)
    - Section titles as H2 (## heading) - only sections with questions
    - Bullet points (- ) for each question/task

    Args:
        lab: The Lab object containing all lab data.
        output_path: The path where the document should be saved.

    Returns:
        The output path where the document was saved.
    """
    lines: list[str] = []

    # Add lab title as H1
    title = f"{lab.number}: {lab.title}"
    lines.append(f"# {title}")
    lines.append("")

    # Process each section (only include sections with questions)
    for section in lab.sections:
        # Skip sections with no questions
        if not section.questions:
            continue

        # Add section title as H2
        section_title = f"{section.number}. {section.title}"
        lines.append(f"## {section_title}")
        lines.append("")

        # Add bullet points for each question
        for question in section.questions:
            lines.append(f"- {question.text}")

        lines.append("")

    # Ensure output directory exists
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    # Write the markdown file
    content = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    return output_path
